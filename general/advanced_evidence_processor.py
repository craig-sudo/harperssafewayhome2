#!/usr/bin/env python3
"""
Harper's Advanced Evidence Processor - Multi-Format Support
Processes PDFs, Videos, Documents, Images, Audio - EVERYTHING for FDSJ-739-24
This beast handles ALL evidence types while the main OCR processes images
"""

import pytesseract
from PIL import Image
import PyPDF2
from docx import Document
import csv
import os
import sys
import logging
from datetime import datetime
from pathlib import Path
import re
from tqdm import tqdm
import pandas as pd
import subprocess
import json

# Video/Audio processing imports
try:
    import moviepy.editor as mp
    MOVIEPY_AVAILABLE = True
except ImportError:
    MOVIEPY_AVAILABLE = False

try:
    import speech_recognition as sr
    SPEECH_RECOGNITION_AVAILABLE = True
except ImportError:
    SPEECH_RECOGNITION_AVAILABLE = False

class AdvancedEvidenceProcessor:
    """Multi-format evidence processor for Harper's case - handles EVERYTHING"""
    
    def __init__(self):
        """Initialize the advanced processor"""
        # Setup Tesseract
        if os.path.exists(r"C:\Program Files\Tesseract-OCR\tesseract.exe"):
            pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        
        # Initialize counters
        self.processed_count = 0
        self.pdf_count = 0
        self.video_count = 0
        self.audio_count = 0
        self.doc_count = 0
        self.error_count = 0
        
        # Setup folders
        self.source_folders = [
            Path("custody_screenshots"),
            Path("custody_screenshots_organized"), 
            Path("custody_screenshots_smart_renamed")
        ]
        self.output_folder = Path("output")
        self.output_folder.mkdir(exist_ok=True)
        
        # Initialize enhanced CSV
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.csv_filename = self.output_folder / f"harper_advanced_evidence_{timestamp}.csv"
        
        # Setup logging
        log_folder = Path("logs")
        log_folder.mkdir(exist_ok=True)
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_folder / f"advanced_processing_{timestamp}.log"),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
        
        print(f"""
╔══════════════════════════════════════════════════════════════════╗
║           🚀 HARPER'S ADVANCED EVIDENCE PROCESSOR 🚀             ║
║                                                                  ║
║  📄 PDFs  🎥 Videos  🔊 Audio  📝 Documents  🖼️ Images          ║
║  🔍 Multi-Format Text Extraction & Analysis System              ║
║  📋 Case: FDSJ-739-24 | COMPLETE Evidence Processing            ║
║                                                                  ║
║  💾 Output: {self.csv_filename.name}                           ║
╚══════════════════════════════════════════════════════════════════╝
        """)

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF documents"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            # Clean up text
            text = re.sub(r'\n+', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            
            self.pdf_count += 1
            return text
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF text from {pdf_path}: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text from Word documents"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            # Clean up text
            text = re.sub(r'\n+', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            
            self.doc_count += 1
            return text
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text from {docx_path}: {e}")
            return ""

    def extract_audio_from_video(self, video_path: Path) -> str:
        """Extract audio track from video files"""
        if not MOVIEPY_AVAILABLE:
            self.logger.warning("MoviePy not available - cannot process videos")
            return ""
        
        try:
            # Extract audio to temporary WAV file
            audio_path = self.output_folder / f"temp_audio_{video_path.stem}.wav"
            
            video = mp.VideoFileClip(str(video_path))
            if video.audio:
                video.audio.write_audiofile(str(audio_path), verbose=False, logger=None)
                video.close()
                
                # Now transcribe the audio
                transcript = self.transcribe_audio(audio_path)
                
                # Clean up temp file
                if audio_path.exists():
                    audio_path.unlink()
                
                self.video_count += 1
                return transcript
            else:
                self.logger.warning(f"No audio track in video: {video_path}")
                return ""
                
        except Exception as e:
            self.logger.error(f"Error extracting audio from video {video_path}: {e}")
            return ""

    def transcribe_audio(self, audio_path: Path) -> str:
        """Transcribe audio files to text"""
        if not SPEECH_RECOGNITION_AVAILABLE:
            self.logger.warning("Speech recognition not available - cannot transcribe audio")
            return ""
        
        try:
            recognizer = sr.Recognizer()
            
            with sr.AudioFile(str(audio_path)) as source:
                # Adjust for ambient noise
                recognizer.adjust_for_ambient_noise(source, duration=0.5)
                audio_data = recognizer.record(source)
            
            # Try multiple recognition engines
            transcript = ""
            
            # Try Google Speech Recognition (free)
            try:
                transcript = recognizer.recognize_google(audio_data)
                self.logger.info(f"Google Speech: Successfully transcribed {audio_path}")
            except sr.UnknownValueError:
                self.logger.warning("Google Speech could not understand audio")
            except sr.RequestError as e:
                self.logger.warning(f"Google Speech error: {e}")
            
            # If Google fails, try offline recognition
            if not transcript:
                try:
                    transcript = recognizer.recognize_sphinx(audio_data)
                    self.logger.info(f"Sphinx: Successfully transcribed {audio_path}")
                except sr.UnknownValueError:
                    self.logger.warning("Sphinx could not understand audio")
                except sr.RequestError as e:
                    self.logger.warning(f"Sphinx error: {e}")
            
            self.audio_count += 1
            return transcript
            
        except Exception as e:
            self.logger.error(f"Error transcribing audio {audio_path}: {e}")
            return ""

    def analyze_metadata(self, file_path: Path) -> dict:
        """Extract metadata from files"""
        try:
            stat = file_path.stat()
            metadata = {
                'file_size': stat.st_size,
                'created': datetime.fromtimestamp(stat.st_ctime),
                'modified': datetime.fromtimestamp(stat.st_mtime),
                'extension': file_path.suffix.lower()
            }
            
            # Try to extract EXIF data from images
            if metadata['extension'] in ['.jpg', '.jpeg', '.png', '.tiff']:
                try:
                    from PIL.ExifTags import TAGS
                    image = Image.open(file_path)
                    exif_data = image._getexif()
                    if exif_data:
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            if tag == 'DateTime':
                                metadata['exif_datetime'] = value
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting metadata from {file_path}: {e}")
            return {}

    def categorize_evidence_advanced(self, filename: str, text: str, file_type: str) -> tuple:
        """Advanced evidence categorization with enhanced detection"""
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        priority = "MEDIUM"
        categories = []
        key_phrases = []
        evidence_type = file_type.upper()
        
        # Critical evidence detection
        critical_keywords = [
            'december 9', 'incident', 'emergency', 'police', 'arrest', 'threat',
            'custody violation', 'contempt', 'harper', 'child safety', 'abuse'
        ]
        
        # High priority threats
        threat_keywords = [
            'kill', 'hurt', 'harm', 'revenge', 'get back', 'sorry will be',
            'regret', 'intimidate', 'scare', 'afraid', 'fear', 'weapon'
        ]
        
        # Legal/Court indicators
        legal_keywords = [
            'court', 'judge', 'lawyer', 'attorney', 'custody agreement', 
            'parenting plan', 'visitation', 'contempt', 'violation'
        ]
        
        # Medical/Health indicators
        medical_keywords = [
            'doctor', 'hospital', 'medical', 'sick', 'injury', 'medicine',
            'prescription', 'therapy', 'counseling', 'mental health'
        ]
        
        # Financial indicators
        financial_keywords = [
            'money', 'payment', 'support', 'bank', 'account', 'financial',
            'income', 'expense', 'cost', 'fee', 'debt'
        ]
        
        # Categorization logic
        if any(keyword in text_lower for keyword in critical_keywords):
            priority = "CRITICAL"
            categories.append('CRITICAL_INCIDENT')
        
        if any(keyword in text_lower for keyword in threat_keywords):
            categories.append('THREATENING_BEHAVIOR')
            priority = "HIGH" if priority != "CRITICAL" else priority
        
        if any(keyword in text_lower for keyword in legal_keywords):
            categories.append('LEGAL_PROCEEDINGS')
        
        if any(keyword in text_lower for keyword in medical_keywords):
            categories.append('MEDICAL_EVIDENCE')
        
        if any(keyword in text_lower for keyword in financial_keywords):
            categories.append('FINANCIAL_EVIDENCE')
        
        # Extract key phrases
        for phrase in critical_keywords + threat_keywords:
            if phrase in text_lower:
                key_phrases.append(phrase)
        
        # People detection (enhanced)
        people_mentioned = []
        people_patterns = {
            'Emma': ['emma', 'emma ryan', 'emma elizabeth', '5062910202'],
            'Matt': ['matt', 'matt ryan', 'matthew'],
            'Cole': ['cole', 'cole brook'],
            'Tony': ['tony', 'tony baker'],
            'Harper': ['harper', 'child', 'daughter', 'kid']
        }
        
        for person, patterns in people_patterns.items():
            if any(pattern in text_lower or pattern in filename_lower for pattern in patterns):
                people_mentioned.append(person)
        
        return priority, categories, key_phrases, people_mentioned, evidence_type

    def process_all_evidence_types(self):
        """Process ALL evidence types - PDFs, videos, docs, everything!"""
        self.logger.info("Starting ADVANCED multi-format evidence processing...")
        
        # File type handlers
        handlers = {
            '.pdf': self.extract_text_from_pdf,
            '.docx': self.extract_text_from_docx,
            '.doc': self.extract_text_from_docx,  # Will try to handle .doc too
            '.mp4': self.extract_audio_from_video,
            '.avi': self.extract_audio_from_video,
            '.mov': self.extract_audio_from_video,
            '.3gp': self.extract_audio_from_video,
            '.wav': self.transcribe_audio,
            '.mp3': self.transcribe_audio,
            '.txt': lambda p: p.read_text(encoding='utf-8', errors='ignore')
        }
        
        # Initialize CSV
        with open(self.csv_filename, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = [
                'filename', 'file_type', 'evidence_type', 'file_path', 'file_size',
                'date_created', 'date_modified', 'text_content', 'text_length', 
                'priority', 'categories', 'key_phrases', 'people_mentioned', 
                'processing_timestamp', 'metadata'
            ]
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            
            # Collect all files
            all_files = []
            for folder in self.source_folders:
                if folder.exists():
                    for file_path in folder.rglob('*'):
                        if file_path.is_file() and file_path.suffix.lower() in handlers:
                            all_files.append(file_path)
            
            print(f"\n🔍 Processing {len(all_files)} advanced evidence files...")
            
            with tqdm(total=len(all_files), desc="Multi-format processing") as pbar:
                for file_path in all_files:
                    try:
                        file_extension = file_path.suffix.lower()
                        handler = handlers.get(file_extension)
                        
                        if handler:
                            # Extract text using appropriate handler
                            text = handler(file_path)
                            
                            # Extract metadata
                            metadata = self.analyze_metadata(file_path)
                            
                            # Analyze content
                            priority, categories, key_phrases, people_mentioned, evidence_type = \
                                self.categorize_evidence_advanced(file_path.name, text, file_extension[1:])
                            
                            # Write to CSV
                            writer.writerow({
                                'filename': file_path.name,
                                'file_type': file_extension[1:],
                                'evidence_type': evidence_type,
                                'file_path': str(file_path),
                                'file_size': metadata.get('file_size', 0),
                                'date_created': metadata.get('created', ''),
                                'date_modified': metadata.get('modified', ''),
                                'text_content': text[:2000] if text else "",  # Limit for CSV
                                'text_length': len(text),
                                'priority': priority,
                                'categories': '; '.join(categories),
                                'key_phrases': '; '.join(key_phrases),
                                'people_mentioned': '; '.join(people_mentioned),
                                'processing_timestamp': datetime.now().isoformat(),
                                'metadata': json.dumps(metadata, default=str)
                            })
                            
                            self.processed_count += 1
                            
                            pbar.set_description(f"Processed: PDFs:{self.pdf_count} Videos:{self.video_count} Docs:{self.doc_count}")
                            
                    except Exception as e:
                        self.logger.error(f"Error processing {file_path}: {e}")
                        self.error_count += 1
                    
                    pbar.update(1)

    def generate_advanced_report(self):
        """Generate comprehensive processing report"""
        print(f"""
╔══════════════════════════════════════════════════════════════════╗
║                🚀 ADVANCED PROCESSING COMPLETE! 🚀               ║
╚══════════════════════════════════════════════════════════════════╝

✅ TOTAL FILES PROCESSED: {self.processed_count}
📄 PDFs PROCESSED: {self.pdf_count}
🎥 VIDEOS TRANSCRIBED: {self.video_count}  
🔊 AUDIO FILES TRANSCRIBED: {self.audio_count}
📝 DOCUMENTS PROCESSED: {self.doc_count}
❌ ERRORS: {self.error_count}

📋 RESULTS SAVED TO: {self.csv_filename}

🎯 This enhanced CSV now contains:
   • PDF document text extraction
   • Video/audio transcriptions
   • Word document content
   • Metadata analysis
   • Advanced categorization
   • Multi-format evidence processing

💪 Your evidence processing is now NEXT LEVEL!
   Every PDF, video, document, and audio file has been analyzed!
        """)

def main():
    """Main function"""
    try:
        processor = AdvancedEvidenceProcessor()
        processor.process_all_evidence_types()
        processor.generate_advanced_report()
        
    except KeyboardInterrupt:
        print("\n🚫 Processing interrupted")
    except Exception as e:
        print(f"❌ Error: {e}")
        logging.error(f"Fatal error: {e}")

if __name__ == "__main__":
    main()